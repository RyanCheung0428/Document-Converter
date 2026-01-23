"""
Document format converter
Handles conversions between Word, Excel, PDF, and plain text formats
Pure Python implementation - no external dependencies
"""
import os
import csv
import shutil
import tempfile
from pathlib import Path
from docx import Document
from openpyxl import load_workbook, Workbook
from PyPDF2 import PdfReader, PdfWriter
from pdf2docx import Converter as PDFConverter
from PIL import Image
import img2pdf
from reportlab.lib.pagesizes import letter, A4
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
import io


class DocumentConverter:
    """Converts between document formats (Word, Excel, PDF, TXT, MD, CSV)"""
    
    # Word-like formats (pure Python support)
    WORD_FORMATS = {'docx', 'txt', 'md'}
    
    # Excel-like formats (pure Python support)
    EXCEL_FORMATS = {'xlsx', 'xlsm', 'csv'}
    
    def __init__(self):
        """Initialize document converter"""
        pass
    
    def convert(self, input_path, output_path, target_format):
        """
        Convert document to target format
        
        Args:
            input_path: Source file path
            output_path: Output file path
            target_format: Target format (pdf, docx, xlsx, txt, md, csv, etc.)
            
        Returns:
            bool: True if successful
        """
        source_ext = os.path.splitext(input_path)[1].lower().lstrip('.')
        target_format = target_format.lower()
        
        # Same format - just copy
        if source_ext == target_format:
            shutil.copy2(input_path, output_path)
            return True
        
        # Route to appropriate converter
        if source_ext == 'pdf':
            return self._convert_from_pdf(input_path, output_path, target_format)
        elif source_ext == 'docx':
            return self._convert_from_docx(input_path, output_path, target_format)
        elif source_ext in ['txt', 'md']:
            return self._convert_from_text(input_path, output_path, source_ext, target_format)
        elif source_ext in ['xlsx', 'xlsm']:
            return self._convert_from_excel(input_path, output_path, source_ext, target_format)
        elif source_ext == 'csv':
            return self._convert_from_csv(input_path, output_path, target_format)
        else:
            raise ValueError(f"不支援的來源格式: {source_ext}")
    
    def _convert_from_pdf(self, input_path, output_path, target_format):
        """Convert from PDF to other formats"""
        if target_format == 'docx':
            return self._pdf_to_word(input_path, output_path)
        elif target_format == 'txt':
            return self._pdf_to_text(input_path, output_path)
        elif target_format in ['png', 'jpg', 'jpeg']:
            return self._pdf_to_image(input_path, output_path, target_format)
        else:
            raise ValueError(f"不支援從 PDF 轉換至 {target_format}")
    
    def _convert_from_docx(self, input_path, output_path, target_format):
        """Convert from DOCX to other formats"""
        if target_format == 'pdf':
            return self._word_to_pdf(input_path, output_path)
        elif target_format == 'txt':
            return self._docx_to_text(input_path, output_path)
        elif target_format == 'md':
            return self._docx_to_markdown(input_path, output_path)
        elif target_format == 'docx':
            shutil.copy2(input_path, output_path)
            return True
        else:
            raise ValueError(f"不支援從 DOCX 轉換至 {target_format}")
    
    def _convert_from_text(self, input_path, output_path, source_ext, target_format):
        """Convert from TXT/MD to other formats"""
        if target_format == 'pdf':
            return self._text_to_pdf(input_path, output_path)
        elif target_format == 'docx':
            return self._text_to_docx(input_path, output_path)
        elif target_format in ['txt', 'md']:
            # txt <-> md is essentially a copy (content is the same, extension differs)
            shutil.copy2(input_path, output_path)
            return True
        else:
            raise ValueError(f"不支援從 {source_ext.upper()} 轉換至 {target_format}")
    
    def _convert_from_excel(self, input_path, output_path, source_ext, target_format):
        """Convert from Excel (xlsx/xlsm) to other formats"""
        if target_format == 'pdf':
            return self._excel_to_pdf(input_path, output_path)
        elif target_format == 'csv':
            return self._excel_to_csv(input_path, output_path)
        elif target_format == 'xlsx':
            if source_ext == 'xlsm':
                # xlsm -> xlsx: save without macros
                wb = load_workbook(input_path, keep_vba=False)
                wb.save(output_path)
                return True
            else:
                shutil.copy2(input_path, output_path)
                return True
        elif target_format == 'xlsm':
            if source_ext == 'xlsm':
                shutil.copy2(input_path, output_path)
                return True
            else:
                # xlsx -> xlsm: cannot add macros, just save as xlsm extension
                # Note: This won't add macros, just changes extension
                wb = load_workbook(input_path)
                wb.save(output_path)
                return True
        else:
            raise ValueError(f"不支援從 Excel 轉換至 {target_format}")
    
    def _convert_from_csv(self, input_path, output_path, target_format):
        """Convert from CSV to other formats"""
        if target_format == 'pdf':
            # CSV -> xlsx -> PDF
            with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp:
                tmp_xlsx = tmp.name
            try:
                self._csv_to_excel(input_path, tmp_xlsx)
                return self._excel_to_pdf(tmp_xlsx, output_path)
            finally:
                if os.path.exists(tmp_xlsx):
                    os.unlink(tmp_xlsx)
        elif target_format == 'xlsx':
            return self._csv_to_excel(input_path, output_path)
        elif target_format == 'csv':
            shutil.copy2(input_path, output_path)
            return True
        else:
            raise ValueError(f"不支援從 CSV 轉換至 {target_format}")
    
    # ==================== Text/Markdown conversions ====================
    
    def _docx_to_text(self, input_path, output_path):
        """Extract plain text from DOCX"""
        try:
            doc = Document(input_path)
            lines = []
            
            for para in doc.paragraphs:
                lines.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = '\t'.join(cell.text.strip() for cell in row.cells)
                    lines.append(row_text)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(lines))
            
            return True
        except Exception as e:
            raise Exception(f"DOCX 轉文字失敗: {str(e)}")
    
    def _docx_to_markdown(self, input_path, output_path):
        """Convert DOCX to Markdown (basic conversion)"""
        try:
            doc = Document(input_path)
            lines = []
            
            for para in doc.paragraphs:
                text = para.text
                if not text.strip():
                    lines.append('')
                    continue
                
                # Handle headings
                style_name = para.style.name if para.style else ''
                if style_name and 'Heading' in style_name:
                    # Extract heading level
                    level = 1
                    for char in style_name:
                        if char.isdigit():
                            level = int(char)
                            break
                    lines.append('#' * level + ' ' + text)
                else:
                    # Check for bold/italic in runs
                    md_text = ''
                    for run in para.runs:
                        run_text = run.text
                        if run.bold and run.italic:
                            md_text += f'***{run_text}***'
                        elif run.bold:
                            md_text += f'**{run_text}**'
                        elif run.italic:
                            md_text += f'*{run_text}*'
                        else:
                            md_text += run_text
                    lines.append(md_text if md_text else text)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n\n'.join(lines))
            
            return True
        except Exception as e:
            raise Exception(f"DOCX 轉 Markdown 失敗: {str(e)}")
    
    def _text_to_docx(self, input_path, output_path):
        """Convert plain text or Markdown to DOCX"""
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            doc = Document()
            
            # Simple markdown parsing
            is_markdown = input_path.lower().endswith('.md')
            
            for line in content.split('\n'):
                if not line.strip():
                    doc.add_paragraph('')
                    continue
                
                if is_markdown:
                    # Handle markdown headings
                    if line.startswith('######'):
                        doc.add_heading(line[6:].strip(), level=6)
                    elif line.startswith('#####'):
                        doc.add_heading(line[5:].strip(), level=5)
                    elif line.startswith('####'):
                        doc.add_heading(line[4:].strip(), level=4)
                    elif line.startswith('###'):
                        doc.add_heading(line[3:].strip(), level=3)
                    elif line.startswith('##'):
                        doc.add_heading(line[2:].strip(), level=2)
                    elif line.startswith('#'):
                        doc.add_heading(line[1:].strip(), level=1)
                    elif line.startswith('- ') or line.startswith('* '):
                        # List item
                        doc.add_paragraph(line[2:], style='List Bullet')
                    elif line.startswith('```'):
                        # Skip code block markers
                        continue
                    else:
                        # Regular paragraph - handle bold/italic
                        para = doc.add_paragraph()
                        self._parse_md_inline(para, line)
                else:
                    doc.add_paragraph(line)
            
            doc.save(output_path)
            return True
        except Exception as e:
            raise Exception(f"文字轉 DOCX 失敗: {str(e)}")
    
    def _parse_md_inline(self, para, text):
        """Parse markdown inline formatting (bold, italic) into docx runs"""
        import re
        
        # Simple approach: just add plain text for now
        # More sophisticated parsing would require proper tokenization
        remaining = text
        
        # Try to find bold text
        bold_pattern = re.compile(r'\*\*(.+?)\*\*')
        italic_pattern = re.compile(r'\*([^*]+?)\*')
        
        pos = 0
        for match in bold_pattern.finditer(text):
            # Add text before match
            if match.start() > pos:
                para.add_run(text[pos:match.start()])
            # Add bold text
            run = para.add_run(match.group(1))
            run.bold = True
            pos = match.end()
        
        # Add remaining text
        if pos < len(text):
            para.add_run(text[pos:])
    
    def _text_to_pdf(self, input_path, output_path):
        """Convert text/markdown to PDF"""
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            doc = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            styles = getSampleStyleSheet()
            story = []
            
            # Create heading styles
            heading_styles = {}
            for i in range(1, 7):
                heading_styles[i] = ParagraphStyle(
                    f'MDHeading{i}',
                    parent=styles['Heading1'],
                    fontSize=24 - (i * 2),
                    spaceAfter=12,
                    spaceBefore=12
                )
            
            is_markdown = input_path.lower().endswith('.md')
            
            for line in content.split('\n'):
                if not line.strip():
                    story.append(Spacer(1, 0.1 * inch))
                    continue
                
                # Escape XML characters
                safe_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                
                if is_markdown:
                    # Handle headings
                    if line.startswith('#'):
                        level = 0
                        for char in line:
                            if char == '#':
                                level += 1
                            else:
                                break
                        level = min(level, 6)
                        text = safe_line[level:].strip()
                        story.append(Paragraph(text, heading_styles.get(level, styles['Normal'])))
                    else:
                        story.append(Paragraph(safe_line, styles['Normal']))
                else:
                    story.append(Paragraph(safe_line, styles['Normal']))
                
                story.append(Spacer(1, 0.05 * inch))
            
            doc.build(story)
            return True
        except Exception as e:
            raise Exception(f"文字轉 PDF 失敗: {str(e)}")
    
    def _pdf_to_text(self, input_path, output_path):
        """Extract text from PDF"""
        try:
            reader = PdfReader(input_path)
            text_content = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_content.append(text)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n\n'.join(text_content))
            
            return True
        except Exception as e:
            raise Exception(f"PDF 轉文字失敗: {str(e)}")
    
    # ==================== CSV/Excel conversions ====================
    
    def _excel_to_csv(self, input_path, output_path):
        """Convert Excel to CSV (active sheet only)"""
        try:
            # Check if xlsm and load with keep_vba
            is_xlsm = input_path.lower().endswith('.xlsm')
            wb = load_workbook(input_path, keep_vba=is_xlsm, data_only=True)
            ws = wb.active
            
            with open(output_path, 'w', newline='', encoding='utf-8-sig') as f:
                writer = csv.writer(f)
                for row in ws.iter_rows(values_only=True):
                    # Convert None to empty string
                    row_data = ['' if cell is None else str(cell) for cell in row]
                    writer.writerow(row_data)
            
            return True
        except Exception as e:
            raise Exception(f"Excel 轉 CSV 失敗: {str(e)}")
    
    def _csv_to_excel(self, input_path, output_path):
        """Convert CSV to Excel"""
        try:
            wb = Workbook()
            ws = wb.active
            
            # Try different encodings
            encodings = ['utf-8-sig', 'utf-8', 'gbk', 'big5', 'latin1']
            content = None
            
            for encoding in encodings:
                try:
                    with open(input_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise Exception("無法識別 CSV 檔案編碼")
            
            # Parse CSV
            reader = csv.reader(content.splitlines())
            for row_idx, row in enumerate(reader, 1):
                for col_idx, value in enumerate(row, 1):
                    # Try to convert to number if possible
                    try:
                        if '.' in value:
                            ws.cell(row=row_idx, column=col_idx, value=float(value))
                        else:
                            ws.cell(row=row_idx, column=col_idx, value=int(value))
                    except (ValueError, AttributeError):
                        ws.cell(row=row_idx, column=col_idx, value=value)
            
            wb.save(output_path)
            return True
        except Exception as e:
            raise Exception(f"CSV 轉 Excel 失敗: {str(e)}")
    
    # ==================== Existing conversion methods ====================
    
    def _pdf_to_word(self, input_path, output_path):
        """Convert PDF to Word document"""
        try:
            cv = PDFConverter(input_path)
            cv.convert(output_path)
            cv.close()
            return True
        except Exception as e:
            raise Exception(f"PDF 轉 Word 失敗: {str(e)}")
    
    def _pdf_to_image(self, input_path, output_path, image_format):
        """Convert first page of PDF to image"""
        try:
            import fitz  # PyMuPDF
            
            pdf_document = fitz.open(input_path)
            page = pdf_document[0]
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            
            if image_format in ['jpg', 'jpeg']:
                pix.save(output_path, 'jpeg')
            else:
                pix.save(output_path, 'png')
            
            pdf_document.close()
            return True
            
        except ImportError:
            raise Exception("PDF 轉圖片需要 PyMuPDF。請執行: pip install PyMuPDF")
        except Exception as e:
            raise Exception(f"PDF 轉圖片失敗: {str(e)}")
    
    def _word_to_pdf(self, input_path, output_path):
        """Convert Word document to PDF"""
        try:
            from docx2pdf import convert
            convert(input_path, output_path)
            return True
        except Exception as e:
            print(f"docx2pdf failed: {e}")
            try:
                return self._word_to_pdf_manual(input_path, output_path)
            except Exception as e2:
                raise Exception(f"Word 轉 PDF 失敗: {str(e2)}")
    
    def _word_to_pdf_manual(self, input_path, output_path):
        """Manual Word to PDF conversion using reportlab"""
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import (
                SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            )
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
            
            doc = Document(input_path)
            
            pdf_doc = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            styles = getSampleStyleSheet()
            story = []
            
            heading_styles = {
                'Heading 1': ParagraphStyle(
                    'CustomHeading1',
                    parent=styles['Heading1'],
                    fontSize=24,
                    textColor=colors.HexColor('#2F5496'),
                    spaceAfter=12,
                    spaceBefore=12,
                    leading=28
                ),
                'Heading 2': ParagraphStyle(
                    'CustomHeading2',
                    parent=styles['Heading2'],
                    fontSize=18,
                    textColor=colors.HexColor('#2F5496'),
                    spaceAfter=10,
                    spaceBefore=10,
                    leading=22
                ),
                'Heading 3': ParagraphStyle(
                    'CustomHeading3',
                    parent=styles['Heading3'],
                    fontSize=14,
                    textColor=colors.HexColor('#2F5496'),
                    spaceAfter=8,
                    spaceBefore=8,
                    leading=18
                ),
            }
            
            for para in doc.paragraphs:
                if not para.text.strip():
                    story.append(Spacer(1, 0.1 * inch))
                    continue
                
                style_name = para.style.name if para.style else 'Normal'
                
                if style_name and style_name.startswith('Heading'):
                    if style_name in heading_styles:
                        style = heading_styles[style_name]
                    elif 'Heading 1' in style_name:
                        style = heading_styles['Heading 1']
                    elif 'Heading 2' in style_name:
                        style = heading_styles['Heading 2']
                    elif 'Heading 3' in style_name:
                        style = heading_styles['Heading 3']
                    else:
                        style = styles['Heading1']
                else:
                    style = styles['Normal']
                
                if para.alignment == 1:
                    style = ParagraphStyle('TempCenter', parent=style, alignment=TA_CENTER)
                elif para.alignment == 2:
                    style = ParagraphStyle('TempRight', parent=style, alignment=TA_RIGHT)
                elif para.alignment == 3:
                    style = ParagraphStyle('TempJustify', parent=style, alignment=TA_JUSTIFY)
                
                formatted_text = ""
                for run in para.runs:
                    text = run.text
                    if not text:
                        continue
                    
                    text = text.replace('&', '&amp;')
                    text = text.replace('<', '&lt;')
                    text = text.replace('>', '&gt;')
                    
                    if run.bold and run.italic:
                        text = f"<b><i>{text}</i></b>"
                    elif run.bold:
                        text = f"<b>{text}</b>"
                    elif run.italic:
                        text = f"<i>{text}</i>"
                    
                    if run.underline:
                        text = f"<u>{text}</u>"
                    
                    formatted_text += text
                
                if formatted_text.strip():
                    try:
                        story.append(Paragraph(formatted_text, style))
                        story.append(Spacer(1, 0.1 * inch))
                    except Exception:
                        story.append(Paragraph(para.text, style))
                        story.append(Spacer(1, 0.1 * inch))
            
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_data.append(cell_text if cell_text else '')
                    table_data.append(row_data)
                
                if table_data:
                    pdf_table = Table(table_data)
                    table_style = TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                        ('FONTSIZE', (0, 1), (-1, -1), 9),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                        ('BOX', (0, 0), (-1, -1), 1, colors.black),
                        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
                        ('TOPPADDING', (0, 0), (-1, -1), 6),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                        ('LEFTPADDING', (0, 0), (-1, -1), 8),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ])
                    pdf_table.setStyle(table_style)
                    story.append(pdf_table)
                    story.append(Spacer(1, 0.2 * inch))
            
            pdf_doc.build(story)
            return True
            
        except Exception as e:
            raise Exception(f"Word 轉 PDF 失敗: {str(e)}")
    
    def _excel_to_pdf(self, input_path, output_path):
        """Convert Excel to PDF with table formatting"""
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4, landscape
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
            from reportlab.lib.units import mm
            
            # Handle xlsm files
            is_xlsm = input_path.lower().endswith('.xlsm')
            wb = load_workbook(input_path, keep_vba=is_xlsm, data_only=True)
            ws = wb.active
            
            data = []
            if ws:
                for row in ws.iter_rows(values_only=True):
                    row_data = [str(cell)[:100] if cell is not None else '' for cell in row]
                    data.append(row_data)
            
            if not data:
                raise Exception("Excel 檔案中沒有資料")
            
            num_cols = len(data[0]) if data else 0
            pagesize = landscape(A4) if num_cols > 6 else A4
            
            doc = SimpleDocTemplate(
                output_path,
                pagesize=pagesize,
                rightMargin=20,
                leftMargin=20,
                topMargin=20,
                bottomMargin=20
            )
            
            page_width = pagesize[0] - 40
            if num_cols > 0:
                col_width = page_width / num_cols
                col_width = max(col_width, 30)
                col_widths = [min(col_width, page_width / num_cols)] * num_cols
            else:
                col_widths = None
            
            table = Table(data, colWidths=col_widths, repeatRows=1)
            
            table_style = TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ])
            
            table.setStyle(table_style)
            doc.build([table])
            
            return True
            
        except Exception as e:
            raise Exception(f"Excel 轉 PDF 失敗: {str(e)}")
    
    def batch_convert(self, input_paths, output_dir, target_format):
        """
        Convert multiple documents
        
        Args:
            input_paths: List of input file paths
            output_dir: Output directory
            target_format: Target format
            
        Returns:
            dict: Results with success/failure for each file
        """
        results = {}
        
        for input_path in input_paths:
            try:
                filename = os.path.splitext(os.path.basename(input_path))[0]
                output_path = os.path.join(output_dir, f"{filename}.{target_format}")
                
                self.convert(input_path, output_path, target_format)
                results[input_path] = {'status': 'success', 'output': output_path}
            except Exception as e:
                results[input_path] = {'status': 'failed', 'error': str(e)}
        
        return results
